from pathlib import Path
import os
import subprocess

def clear():
    os.system("cls" if os.name == "nt" else "clear")

def input_pdf():
    while True:
        raw = input("\nPath PDF file : ").strip().strip('"').strip("'")
        p   = Path(raw)
        if not p.exists():
            print("[!] File not found.")
        elif p.suffix.lower() != ".pdf":
            print("[!] not PDF File.")
        else:
            return p

def input_jpg():
    while True:
        raw = input("Path JPG file : ").strip().strip('"').strip('"')
        p   = Path(raw)
        if not p.exists():
            print("[!] File not found.")
        elif p.suffix.lower() != ".jpg":
            print("[!] not JPG File.")
        else:
            return p

def input_doc():
    while True:
        raw = input("Path DOC file : ").strip().strip('"').strip('"')
        p   = Path(raw)
        if not p.exists():
            print("[!] File not found.")
        elif p.suffix.lower() != ".DOC":
            print("[!] not DOC File.")
        else:
            return p

def pdf_info(pdf_path: Path):
    from pypdf import PdfReader

    read = PdfReader(str(pdf_path))
    meta = read.metadata
    size = pdf_path.stat().st_size / 1024
    page = len(read.pages)

    print(f"File    : {pdf_path.name}")
    print(f"Pages   : {page}")
    print(f"Size    : {size:.1f} KB ({size / 1024:.2f}) MB")
    if meta:
        if meta.title: print(f"Title    : {meta.title}")
        if meta.author: print(f" Author : {meta.author}")

def jpg_info(jpg_path: Path):
    from PIL import Image

    size = jpg_path.stat().st_size / 1024
    img  = Image.open(jpg_path)

    print(f"File    : {jpg_path.name}")
    print(f"Size    : {size:.1f} KB ({size / 1024:.2f}) MB")
    print(f"Dimensi : {img.width} x {img.height} px")
    print(f"Format  : {img.format}")
    print(f"Mode    : {img.mode}")

    exif = img.getexif()
    if exif:
        print(f"EXIF    : {len(exif)} tag found")

    img.close()

def doc_info(doc_path: Path):
    from docx import Document 

    read = Document(str(doc_path))
    meta = read.core_properties
    size = doc_path.stat().st_size / 1024

    n_para = len(read.paragraphs)
    n_word = sum(len(p.text.split()) for p in read.paragraphs)

    print(f"File       : {doc_path.name}")
    print(f"Size       : {size:.1f} KB ({size / 1024:.2f}) MB")
    print(f"Paragraph  : {n_para}")
    print(f"Words      : {n_word}")
    if meta.title:  print(f"Title      : {meta.title}")
    if meta.author: print(f"Author     : {meta.author}")

def pdf_to_jpg(pdf_path: Path, dpi: int = 150, quality: int = 85,
               output_dir: Path = None):
    
    from pdf2image import convert_from_path

    if output_dir is None:
        output_dir = pdf_path.parent 

    print(f"\n[*] Convert PDF → JPG  |  DPI: {dpi}  |  Quality: {quality}%")

    try:
        image = convert_from_path(str(pdf_path), dpi=dpi)
        for i, img in enumerate(image):
            out = output_dir / f"{i + 1:03d}.jpg"
            img.convert("RGB").save(str(out), "JPEG", quality=quality, optimize=True)
            
            print(f"[+] Pages {i + 1:>3} → {out.name}  "
                  f"({out.stat().st_size / 1024:.0f} KB)")
    
        print(f"\n[✓] {len(image)} saved in: {output_dir.resolve()}")

    except Exception as e:
        print(f"\n[ERROR] Failed to Convert: {e}")
        if any(k in str(e).lower() for k in ["poppler", "pdftoppm", "pdfinfo"]):
            print("\n[!] Poppler was not installed in system.")
            print("    Ubuntu/Debian : sudo apt install poppler-utils")
            print("    macOS         : brew install poppler")
            print("    Windows       : https://github.com/oschwartz10612/poppler-windows/releases")

def pdf_to_doc(pdf_path: Path, output_path: Path = None):
    
    from pdf2docx import Converter

    if output_path is None:
        output_path = pdf_path.parent / f"{pdf_path.stem}.docx"

    print(f"\n[*] Convert PDF → DOCX")

    try:
        cv = Converter(str(pdf_path))
        cv.convert(str(output_path))
        cv.close()

        size_kb = output_path.stat().st_size / 1024
        print(f"\n[✓] Converted: {output_path.name}  ({size_kb:.0f} KB)")
        print(f"    Saved in: {output_path.resolve()}")

    except Exception as e:
        print(f"\n[ERROR] Failed to convert: {e}")

def jpg_to_pdf(jpg_path: Path, output_path: Path = None):

    from PIL import Image

    if output_path is None:
        output_path = jpg_path.parent / f"{jpg_path.stem}.pdf"

    print(f"\n[*] Convert JPG → PDF")

    try:
        img = Image.open(jpg_path).convert("RGB")
        img.save(str(output_path), "PDF", resolution=150.0)
        img.close()

        size_kb = output_path.stat().st_size / 1024
        print(f"\n[✓] Converted: {output_path.name}  ({size_kb:.0f} KB)")
        print(f"    Saved in: {output_path.resolve()}")

    except Exception as e:
        print(f"\n[ERROR] Failed to convert: {e}")

def doc_to_pdf(doc_path: Path, output_dir: Path = None):
    if output_dir is None:
        output_dir = doc_path.parent

    print(f"\n[*] Convert DOCX → PDF")

    try:
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf",
             "--outdir", str(output_dir), str(doc_path)],
            capture_output=True, text=True, timeout=60
        )

        output_path = output_dir / f"{doc_path.stem}.pdf"

        if result.returncode != 0 or not output_path.exists():
            print(f"\n[ERROR] Failed to convert: {result.stderr}")
            return

        size_kb = output_path.stat().st_size / 1024
        print(f"\n[✓] Converted: {output_path.name}  ({size_kb:.0f} KB)")
        print(f"    Saved in: {output_path.resolve()}")

    except FileNotFoundError:
        print("\n[ERROR] LibreOffice was not installed in system.")
        print("    run: sudo apt install libreoffice")

    except subprocess.TimeoutExpired:
        print("\n[ERROR] Conversion timeout (file to big/kompleks).")

    except Exception as e:
        print(f"\n[ERROR] Failed to convert: {e}")

_PDF_PRESET = {
    "low":    {"dpi": 72,  "jpg_quality": 40},
    "medium": {"dpi": 120, "jpg_quality": 65},
}

_IMG_QUALITY = {"low": 35, "medium": 60, "high": 80}

def pdf_compres(path: Path):
    level = ask_compress_level()
    if level is None:
        return

    output      = path.parent / f"{path.stem}_compressed.pdf"
    original_kb = path.stat().st_size / 1024

    print(f"\n[*] Compressing {path.name} — level: {level}")

    try:
        if level == "high":
            import pikepdf
            with pikepdf.open(str(path)) as pdf:
                pdf.save(
                    str(output),
                    compress_streams=True,
                    stream_decode_level=pikepdf.StreamDecodeLevel.generalized,
                    recompress_flate=True,
                )
        else:
            from pdf2image import convert_from_path
            preset = _PDF_PRESET[level]
            images = convert_from_path(str(path), dpi=preset["dpi"])
            rgb    = [img.convert("RGB") for img in images]

            save_kwargs = dict(format="PDF", resolution=preset["dpi"],
                                quality=preset["jpg_quality"], optimize=True)
            if len(rgb) == 1:
                rgb[0].save(str(output), **save_kwargs)
            else:
                rgb[0].save(str(output), save_all=True, append_images=rgb[1:], **save_kwargs)

        new_kb    = output.stat().st_size / 1024
        reduction = (original_kb - new_kb) / original_kb * 100 if original_kb else 0

        print(f"\n[✓] Compressed: {output.name}")
        print(f"    Before : {original_kb:.0f} KB")
        print(f"    After  : {new_kb:.0f} KB")
        print(f"    Saved  : {reduction:.0f}%")
        if level != "high":
            print("    [!] Text in this PDF is no longer selectable.")

    except Exception as e:
        print(f"\n[ERROR] Failed to compress: {e}")

def jpg_compres(path: Path):
    level = ask_compress_level()
    if level is None:
        return

    from PIL import Image

    output      = path.parent / f"{path.stem}_compressed.jpg"
    original_kb = path.stat().st_size / 1024
    quality     = _IMG_QUALITY[level]

    print(f"\n[*] Compressing {path.name} — level: {level}")

    try:
        img = Image.open(path).convert("RGB")
        img.save(str(output), "JPEG", quality=quality, optimize=True)
        img.close()

        new_kb    = output.stat().st_size / 1024
        reduction = (original_kb - new_kb) / original_kb * 100 if original_kb else 0

        print(f"\n[✓] Compressed: {output.name}")
        print(f"    Before : {original_kb:.0f} KB")
        print(f"    After  : {new_kb:.0f} KB")
        print(f"    Saved  : {reduction:.0f}%")

    except Exception as e:
        print(f"\n[ERROR] Failed to compress: {e}")

def doc_compres(path: Path):
    level = ask_compress_level()
    if level is None:
        return

    if path.suffix.lower() != ".docx":
        print("\n[!] Only .docx is supported.")
        input("\nEnter to continue")
        return

    import zipfile, io
    from PIL import Image

    output      = path.parent / f"{path.stem}_compressed.docx"
    original_kb = path.stat().st_size / 1024
    quality     = _IMG_QUALITY[level]

    print(f"\n[*] Compressing {path.name} — level: {level}")

    try:
        with zipfile.ZipFile(str(path), "r") as zin:
            with zipfile.ZipFile(str(output), "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.namelist():
                    data = zin.read(item)

                    is_image = item.startswith("word/media/") and \
                               item.lower().endswith((".jpg", ".jpeg", ".png"))

                    if is_image:
                        try:
                            img = Image.open(io.BytesIO(data)).convert("RGB")
                            buf = io.BytesIO()
                            img.save(buf, "JPEG", quality=quality, optimize=True)
                            data = buf.getvalue()
                        except Exception:
                            pass  # gambar gagal diproses, pakai data asli

                    zout.writestr(item, data)

        new_kb    = output.stat().st_size / 1024
        reduction = (original_kb - new_kb) / original_kb * 100 if original_kb else 0

        print(f"\n[✓] Compressed: {output.name}")
        print(f"    Before : {original_kb:.0f} KB")
        print(f"    After  : {new_kb:.0f} KB")
        print(f"    Saved  : {reduction:.0f}%")

    except Exception as e:
        print(f"\n[ERROR] Failed to compress: {e}")

def ask_compress_level() -> str | None:
    while True:
        clear()
        print("\nCompression level\n")
        print("1\tLow    → smallest size, quality drops noticeably")
        print("2\tMedium → balanced [recommended]")
        print("3\tHigh   → best quality, smaller size reduction")
        print("\n0\tBack")

        tipe = input("\nInput [1-3/0] : ").strip()

        if tipe == "0":
            return None

        mapping = {"1": "low", "2": "medium", "3": "high"}
        if tipe not in mapping:
            print("\n[!] option not valid.")
            input("\nEnter to continue")
            continue

        return mapping[tipe]

def welcome():
    clear()
    print("\n" + "=" * 40)
    print("\t" + "  PDF Tools File")
    print("=" * 40)
    print("1\t→\tInfo File")
    print("2\t→\tConvert File")
    print("3\t→\tCompress File")
    print("\nQ → Quit")

def flow_info_file():
    while True:
        clear()
        print("\nChose your file type\n")
        print("1\tPDF")
        print("2\tJPG")
        print("3\tDOC/DOCX")
        print("\n0\tBack")

        tipe = input("\nInput [1-3/0] : ").strip()

        if tipe == "0":
            return   
        
        if tipe not in ("1", "2", "3"):
            print("\n[!] option not valid.")
            input("\nEnter to continue")
            continue

        path_raw = input("Path file   : ").strip().strip('"').strip("'")
        path = Path(path_raw)

        if not path.exists():
            print("\n[!] File not found.")
            input("\nEnter to cotinue")
            continue   

        clear()
        if   tipe == "1": pdf_info(path)
        elif tipe == "2": jpg_info(path)
        elif tipe == "3": doc_info(path)
        
        input("\nEnter to continue")
            
def flow_convert_file():
    while True:
        clear()
        print("\nwhat do you to convert to?\n")
        print("1\tPDF to JPG")
        print("2\tPDF to DOC/DOCX")
        print("3\tJPG to PDF")
        print("4\tDOC/DOCX to PDF")
        print("\n0\tBack")
    
        tipe = input("\nInput[1-4/0] : ").strip()
        
        if tipe == "0":
            return
        
        if tipe not in ("1", "2", "3", "4"):
            print("\n[!] option not valid.")
            input("\nEnter to continue")
            continue
        
        path_raw = input("Path file   : ").strip().strip('"').strip("'")
        path = Path(path_raw)
    
        if not path.exists():
            print("\n[!] File not found.")
            input("\nEnter to continue")
            continue
    
        clear()
        if   tipe == "1" : pdf_to_jpg(path)
        elif tipe == "2" : pdf_to_doc(path)
        elif tipe == "3" : jpg_to_pdf(path)
        elif tipe == "4" : doc_to_pdf(path)

        input("\nEnter to continue")

def flow_compres_file():
    while True:
        clear()
        print("\nWhat file do you to compres ?\n")
        print("1\tPDF")
        print("2\tJPG")
        print("3\tDOC/DOCX")
        print("\n0\tBack")

        tipe = input("\nInput[1-3/0] : ").strip()

        if tipe == "0":
            return
        
        if tipe not in ("1", "2", "3"):
            print("\n[!] option not valid.")
            input("\nEnter to continue")
            continue

        path_raw = input("Path file : ").strip().strip('"').strip("'")
        path = Path(path_raw)

        if not path.exists():
            print("\n[!] File not found.")
            input("\nEnter to continue")
            continue
        
        clear()
        if   tipe == "1" : pdf_compres(path)
        elif tipe == "2" : jpg_compres(path)
        elif tipe == "3" : doc_compres(path)

        input("\nEnter to continue")


def main():
    while True:
        welcome()
        opsi = input("\nInput [1-3/Q] : ").strip().upper()

        if opsi not in ("1", "2", "3", "Q", "q"):
            print("\n[!] Fucking invalid.")
            input("\nEnter to continue")
            continue
    
        if   opsi == "Q": print("\nQuit"); break
        elif opsi == "1": flow_info_file()
        elif opsi == "2": flow_convert_file()
        elif opsi == "3": flow_compres_file()

    return

if __name__ == "__main__":
    main()
