import copy
from pathlib import Path


def merge_pdf(paths: list[Path], output: Path = None):
    from pypdf import PdfWriter

    if output is None:
        output = paths[0].parent / "merged.pdf"

    print(f"\n[*] Merge {len(paths)} PDF → {output.name}\n")

    writer = PdfWriter()
    try:
        for i, p in enumerate(paths, 1):
            writer.append(str(p))
            print(f"[+] {i}. {p.name}")

        with open(output, "wb") as f:
            writer.write(f)

        size_kb = output.stat().st_size / 1024
        print(f"\n[✓] Merged: {output.name}  ({size_kb:.0f} KB)")
        print(f"    Saved in : {output.resolve()}")

    except Exception as e:
        print(f"\n[ERROR] Merge failed: {e}")

    finally:
        if hasattr(writer, "close"):
            writer.close()

def _parse_ranges(raw: str, total_pages: int):
    raw = raw.strip().lower()
    if raw in ("", "all", "each"):
        return [(i, i) for i in range(1, total_pages + 1)]

    ranges = []
    for part in raw.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            a_str, b_str = part.split("-", 1)
        else:
            a_str = b_str = part
        if not (a_str.isdigit() and b_str.isdigit()):
            return None
        a, b = int(a_str), int(b_str)
        if a < 1 or b > total_pages or a > b:
            return None
        ranges.append((a, b))

    return ranges or None

def split_pdf(path: Path, output_dir: Path = None):
    from pypdf import PdfReader, PdfWriter

    if output_dir is None:
        output_dir = path.parent / f"{path.stem}_split"

    try:
        reader = PdfReader(str(path))
        total = len(reader.pages)
    except Exception as e:
        print(f"\n[ERROR] Unable to open PDF: {e}")
        return

    print(f"\n[*] Split PDF: {path.name}  ({total} pages)")
    print("    input example : 1-3,5,7-9   (blank = each page is a separate file)")
    raw = input("\nPage range:").strip()

    ranges = _parse_ranges(raw, total)
    if ranges is None:
        print(f"\n[!] Invalid range (pages 1-{total}).")
        return

    output_dir.mkdir(parents=True, exist_ok=True)

    try:
        for start, end in ranges:
            writer = PdfWriter()
            for pg in range(start - 1, end):
                writer.add_page(reader.pages[pg])

            name = f"{path.stem}_p{start}.pdf" if start == end else f"{path.stem}_p{start}-{end}.pdf"
            out = output_dir / name

            with open(out, "wb") as f:
                writer.write(f)

            print(f"[+] {name}  (pages {start}-{end})")

        print(f"\n[✓] {len(ranges)} file saved in: {output_dir.resolve()}")

    except Exception as e:
        print(f"\n[ERROR] Failed split: {e}")

def _flatten_ranges(ranges: list[tuple[int, int]]) -> list[int]:
    """[(1,3),(5,5)] -> [1,2,3,5], urut & tanpa duplikat."""
    pages = set()
    for start, end in ranges:
        pages.update(range(start, end + 1))
    return sorted(pages)

def extract_pages(path: Path, output: Path = None):
    """Ambil halaman tertentu, gabung jadi SATU file baru. Beda sama split_pdf
    yang bikin file terpisah per range — ini semua digabung jadi satu output."""
    from pypdf import PdfReader, PdfWriter
 
    try:
        reader = PdfReader(str(path))
        total = len(reader.pages)
    except Exception as e:
        print(f"\n[ERROR] Unable to open PDF: {e}")
        return
 
    print(f"\n[*] Extract pages: {path.name}  ({total} pages)")
    print("    Input example: 1-3,5,7-9   (results are combined into a single file)")
    raw = input("\nPage to be retrieved: ").strip()
 
    ranges = _parse_ranges(raw, total)
    if ranges is None:
        print(f"\n[!] Invalid range (page 1-{total}).")
        return
 
    pages = _flatten_ranges(ranges)
 
    if output is None:
        output = path.parent / f"{path.stem}_extracted.pdf"
 
    try:
        writer = PdfWriter()
        for pg in pages:
            writer.add_page(reader.pages[pg - 1])
 
        with open(output, "wb") as f:
            writer.write(f)
 
        size_kb = output.stat().st_size / 1024
        print(f"\n[✓] Extracted {len(pages)} pages: {output.name}  ({size_kb:.0f} KB)")
        print(f"    Pages  : {', '.join(map(str, pages))}")
        print(f"    Saved in : {output.resolve()}")
 
    except Exception as e:
        print(f"\n[ERROR] Failed extract: {e}")

def delete_pages(path: Path, output: Path = None):
    """Buang halaman tertentu, sisanya digabung jadi satu file — kebalikan
    dari extract_pages."""
    from pypdf import PdfReader, PdfWriter
 
    try:
        reader = PdfReader(str(path))
        total = len(reader.pages)
    except Exception as e:
        print(f"\n[ERROR] Unable to open PDF: {e}")
        return
 
    print(f"\n[*] Delete pages: {path.name}  ({total} pages)")
    print("    Input example: 1-3,5,7-9   (this page will be REMOVED)")
    raw = input("\nPage to be deleted:").strip()
 
    ranges = _parse_ranges(raw, total)
    if ranges is None:
        print(f"\n[!] Invalid range (page 1-{total}).")
        return
 
    to_delete = set(_flatten_ranges(ranges))
    keep = [p for p in range(1, total + 1) if p not in to_delete]
 
    if not keep:
        print("\n[!] You can't delete ALL pages—the PDF can't be empty.")
        return
 
    if output is None:
        output = path.parent / f"{path.stem}_deleted.pdf"
 
    try:
        writer = PdfWriter()
        for pg in keep:
            writer.add_page(reader.pages[pg - 1])
 
        with open(output, "wb") as f:
            writer.write(f)
 
        size_kb = output.stat().st_size / 1024
        print(f"\n[✓] Deleted {len(to_delete)} pages, remaning {len(keep)}: {output.name}  ({size_kb:.0f} KB)")
        print(f"    Saved in : {output.resolve()}")
 
    except Exception as e:
        print(f"\n[ERROR] Failed delete: {e}")

def _ask_rotation_angle() -> int | None:
    print("\nHow many degrees is the rotation (clockwise)?")
    print("1\t90°")
    print("2\t180°")
    print("3\t270°  (= 90° counterclockwise)")
    print("\n0\tBack")
 
    choice = input("\nInput [1-3/0] : ").strip()
    mapping = {"1": 90, "2": 180, "3": 270}
    return mapping.get(choice)

def rotate_pages(path: Path, output: Path = None):
    from pypdf import PdfReader, PdfWriter
 
    try:
        reader = PdfReader(str(path))
        total = len(reader.pages)
    except Exception as e:
        print(f"\n[ERROR] Unable to open PDF: {e}")
        return
 
    print(f"\n[*] Rotate pages: {path.name}  ({total} pages)")
    print("    Input example : 1-3,5,7-9   (blank = all pages)")
    raw = input("\nPages to be rotated:").strip()
 
    ranges = _parse_ranges(raw, total)
    if ranges is None:
        print(f"\n[!] Invalid range (pages 1-{total}).")
        return
 
    to_rotate = set(_flatten_ranges(ranges))
 
    angle = _ask_rotation_angle()
    if angle is None:
        print("\n[!] Canceled.")
        return
 
    if output is None:
        output = path.parent / f"{path.stem}_rotated.pdf"
 
    try:
        writer = PdfWriter()
        for i, page in enumerate(reader.pages, 1):
            if i in to_rotate:
                page.rotate(angle)
            writer.add_page(page)
 
        with open(output, "wb") as f:
            writer.write(f)
 
        size_kb = output.stat().st_size / 1024
        print(f"\n[✓] {len(to_rotate)} rotated page {angle}°: {output.name}  ({size_kb:.0f} KB)")
        print(f"    Saved in : {output.resolve()}")
 
    except Exception as e:
        print(f"\n[ERROR] Failed rotate: {e}")

def merge_docx(paths: list[Path], output: Path = None):
    from docx import Document
    from docxcompose.composer import Composer

    if output is None:
        output = paths[0].parent / "merged.docx"

    print(f"\n[*] Merge {len(paths)} DOCX → {output.name}\n")

    try:
        master = Document(str(paths[0]))
        composer = Composer(master)
        print(f"[+] 1. {paths[0].name}")

        for i, p in enumerate(paths[1:], 2):
            doc = Document(str(p))
            doc.add_page_break()  
            composer.append(doc)
            print(f"[+] {i}. {p.name}")

        composer.save(str(output))

        size_kb = output.stat().st_size / 1024
        print(f"\n[✓] Merged: {output.name}  ({size_kb:.0f} KB)")
        print(f"    Saved in : {output.resolve()}")

    except Exception as e:
        print(f"\n[ERROR] Failed merge: {e}")

def split_docx(path: Path, output_dir: Path = None):
    from docx import Document
    from docx.oxml.ns import qn

    if output_dir is None:
        output_dir = path.parent / f"{path.stem}_split"

    try:
        doc = Document(str(path))
    except Exception as e:
        print(f"\n[ERROR] Unable to open DOCX: {e}")
        return

    print(f"\n[*] Split DOCX: {path.name}")
    print("    [!] Split based on manual page breaks, not rendered pages.\n")

    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))  # section properties (margin, size, dll)

    chunks = [[]]
    has_pagebreak = False
    br_xpath = f".//{qn('w:br')}"

    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        chunks[-1].append(child)

        breaks = child.findall(br_xpath)
        if any(b.get(qn("w:type")) == "page" for b in breaks):
            has_pagebreak = True
            chunks.append([])

    if not has_pagebreak:
        print("[!] There are no manual page breaks in this document.")
        print("    The split was canceled—there was no clear point of division.")
        return

    output_dir.mkdir(parents=True, exist_ok=True)

    try:
        n = 0
        for chunk in chunks:
            if not chunk:
                continue
            n += 1
            new_doc = Document()
            new_body = new_doc.element.body
            for el in list(new_body):
                new_body.remove(el)
            for el in chunk:
                new_body.append(copy.deepcopy(el))
            if sectPr is not None:
                new_body.append(copy.deepcopy(sectPr))

            out = output_dir / f"{path.stem}_part{n}.docx"
            new_doc.save(str(out))
            print(f"[+] {out.name}")

        print(f"\n[✓] {n} file saved in: {output_dir.resolve()}")

    except Exception as e:
        print(f"\n[ERROR] Failed split: {e}")

