import copy
from pathlib import Path


def merge_pdf(paths: list[Path], output: Path = None):
    from pypdf import PdfWriter

    if output is None:
        output = paths[0].parent / "merged.pdf"

    print(f"\n[*] Merge {len(paths)} PDF → {output.name}\n")

    writer = PdfWriter()
    try:
        for i, p in enumerate(paths, 1):
            writer.append(str(p))
            print(f"[+] {i}. {p.name}")

        with open(output, "wb") as f:
            writer.write(f)

        size_kb = output.stat().st_size / 1024
        print(f"\n[✓] Merged: {output.name}  ({size_kb:.0f} KB)")
        print(f"    Saved in : {output.resolve()}")

    except Exception as e:
        print(f"\n[ERROR] Gagal merge: {e}")

    finally:
        if hasattr(writer, "close"):
            writer.close()

def _parse_ranges(raw: str, total_pages: int):
    """'1-3,5,7-9' -> [(1,3),(5,5),(7,9)]. '' / 'all' -> tiap halaman sendiri."""
    raw = raw.strip().lower()
    if raw in ("", "all", "each"):
        return [(i, i) for i in range(1, total_pages + 1)]

    ranges = []
    for part in raw.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            a_str, b_str = part.split("-", 1)
        else:
            a_str = b_str = part
        if not (a_str.isdigit() and b_str.isdigit()):
            return None
        a, b = int(a_str), int(b_str)
        if a < 1 or b > total_pages or a > b:
            return None
        ranges.append((a, b))

    return ranges or None

def split_pdf(path: Path, output_dir: Path = None):
    from pypdf import PdfReader, PdfWriter

    if output_dir is None:
        output_dir = path.parent / f"{path.stem}_split"

    try:
        reader = PdfReader(str(path))
        total = len(reader.pages)
    except Exception as e:
        print(f"\n[ERROR] Gagal baca PDF: {e}")
        return

    print(f"\n[*] Split PDF: {path.name}  ({total} halaman)")
    print("    Contoh input : 1-3,5,7-9   (kosong = tiap halaman jadi 1 file)")
    raw = input("\nRange halaman : ").strip()

    ranges = _parse_ranges(raw, total)
    if ranges is None:
        print(f"\n[!] Range tidak valid (halaman 1-{total}).")
        return

    output_dir.mkdir(parents=True, exist_ok=True)

    try:
        for start, end in ranges:
            writer = PdfWriter()
            for pg in range(start - 1, end):
                writer.add_page(reader.pages[pg])

            name = f"{path.stem}_p{start}.pdf" if start == end else f"{path.stem}_p{start}-{end}.pdf"
            out = output_dir / name

            with open(out, "wb") as f:
                writer.write(f)

            print(f"[+] {name}  (halaman {start}-{end})")

        print(f"\n[✓] {len(ranges)} file saved in: {output_dir.resolve()}")

    except Exception as e:
        print(f"\n[ERROR] Gagal split: {e}")

def merge_docx(paths: list[Path], output: Path = None):
    from docx import Document
    from docxcompose.composer import Composer

    if output is None:
        output = paths[0].parent / "merged.docx"

    print(f"\n[*] Merge {len(paths)} DOCX → {output.name}\n")

    try:
        master = Document(str(paths[0]))
        composer = Composer(master)
        print(f"[+] 1. {paths[0].name}")

        for i, p in enumerate(paths[1:], 2):
            doc = Document(str(p))
            doc.add_page_break()  
            composer.append(doc)
            print(f"[+] {i}. {p.name}")

        composer.save(str(output))

        size_kb = output.stat().st_size / 1024
        print(f"\n[✓] Merged: {output.name}  ({size_kb:.0f} KB)")
        print(f"    Saved in : {output.resolve()}")

    except Exception as e:
        print(f"\n[ERROR] Gagal merge: {e}")

def split_docx(path: Path, output_dir: Path = None):
    from docx import Document
    from docx.oxml.ns import qn

    if output_dir is None:
        output_dir = path.parent / f"{path.stem}_split"

    try:
        doc = Document(str(path))
    except Exception as e:
        print(f"\n[ERROR] Gagal baca DOCX: {e}")
        return

    print(f"\n[*] Split DOCX: {path.name}")
    print("    [!] Split berdasarkan manual page break, bukan halaman hasil render.\n")

    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))  # section properties (margin, size, dll)

    chunks = [[]]
    has_pagebreak = False
    br_xpath = f".//{qn('w:br')}"

    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        chunks[-1].append(child)

        breaks = child.findall(br_xpath)
        if any(b.get(qn("w:type")) == "page" for b in breaks):
            has_pagebreak = True
            chunks.append([])

    if not has_pagebreak:
        print("[!] Nggak ada manual page break di dokumen ini.")
        print("    Split dibatalkan — nggak ada titik potong yang jelas.")
        return

    output_dir.mkdir(parents=True, exist_ok=True)

    try:
        n = 0
        for chunk in chunks:
            if not chunk:
                continue
            n += 1
            new_doc = Document()
            new_body = new_doc.element.body
            for el in list(new_body):
                new_body.remove(el)
            for el in chunk:
                new_body.append(copy.deepcopy(el))
            if sectPr is not None:
                new_body.append(copy.deepcopy(sectPr))

            out = output_dir / f"{path.stem}_part{n}.docx"
            new_doc.save(str(out))
            print(f"[+] {out.name}")

        print(f"\n[✓] {n} file saved in: {output_dir.resolve()}")

    except Exception as e:
        print(f"\n[ERROR] Gagal split: {e}")