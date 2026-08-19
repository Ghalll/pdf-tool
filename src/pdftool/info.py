from pathlib import Path
from .utils import has_javascript, get_attachments, count_images, get_fonts


def jpg_analysis(path: Path):
    from PIL import Image, ExifTags

    try:

        size_kb = path.stat().st_size / 1024
        size_mb = size_kb / 1024

        with Image.open (path) as img:
            exif = img.getexif()
    
            print(f"File    : {path.name}")
            print(f"Size    : {size_kb:.1f} KB  ({size_mb:.2f} MB)")
            print(f"Dimensi : {img.width} x {img.height} px")
            print(f"Format  : {img.format}")
            print(f"Mode    : {img.mode}")

            print("\nEXIF:")

            if not exif:
                print("  None")
                return

            make = exif.get(ExifTags.Base.Make)
            model = exif.get(ExifTags.Base.Model)
            orientation = exif.get(ExifTags.Base.Orientation)
            exif_ifd = exif.get_ifd(ExifTags.IFD.Exif)
            date_taken = exif_ifd.get(ExifTags.Base.DateTimeOriginal)
            gps_ifd = exif.get_ifd(ExifTags.IFD.GPSInfo)

            print(f"  Camera     : {make or '-'}")
            print(f"  Model      : {model or '-'}")
            print(f"  Date Taken : {date_taken or '-'}")
            print(f"  Orientation: {orientation or '-'}")

            if gps_ifd:
                    gps_latitude = gps_ifd.get(ExifTags.GPS.GPSLatitude)
                    gps_latitude_ref = gps_ifd.get(ExifTags.GPS.GPSLatitudeRef)

                    gps_longitude = gps_ifd.get(ExifTags.GPS.GPSLongitude)
                    gps_longitude_ref = gps_ifd.get(ExifTags.GPS.GPSLongitudeRef)

                    if (
                        gps_latitude is not None
                        and gps_latitude_ref is not None
                        and gps_longitude is not None
                        and gps_longitude_ref is not None
                    ):
                        print("  GPS        : Present")

                        def format_gps(coords, ref):
                            try:
                                d = float(coords[0])
                                m = float(coords[1])
                                s = float(coords[2])

                                dec = d + (m / 60.0) + (s / 3600.0)

                                ref_str = str(ref).strip().upper()
                                if ref_str in ['S', 'W']:
                                    dec = -dec
                                return f"{int(d)}° {int(m)}' {s:.3f}\" {ref_str} ({dec:.6f})"
                            except Exception:
                                return f"{coords} {ref}"

                        lat_str = format_gps(gps_latitude, gps_latitude_ref)
                        lon_str = format_gps(gps_longitude, gps_longitude_ref)

                        print(f"    Latitude : {lat_str}")
                        print(f"    Longitude: {lon_str}")
                    else:
                        print("  GPS        : Present, coordinates unavailable")
            else:
                print("  GPS        : None")

    except Exception as e:
        print(f"\n [ERROR] Gagal baca info JPG: {e}")

def doc_info(path: Path, origin: Path | None = None):
    from docx import Document

    origin = origin or path

    try:
        doc    = Document(str(path))
        meta   = doc.core_properties
        size   = path.stat().st_size / 1024
        n_word = sum(len(p.text.split()) for p in doc.paragraphs)
    
        print(f"File       : {path.name}")
        print(f"Size       : {size:.1f} KB  ({size / 1024:.2f} MB)")
        print(f"Paragraphs : {len(doc.paragraphs)}")
        print(f"Words      : {n_word}")
        if meta.title:  print(f"Title      : {meta.title}")
        if meta.author: print(f"Author     : {meta.author}")

    except Exception as e:
        print(f"\n[ERROR] Gagal baca info DOC: {e}")

def pdf_analysis(path: Path):
    from pypdf import PdfReader

    try:
        reader  = PdfReader(str(path))
        meta = reader.metadata
        size_kb = path.stat().st_size / 1024
        size_mb = size_kb / 1024
        
        javascript = has_javascript(reader)
        attachments = get_attachments(reader)
        fonts = get_fonts(reader)
        images = count_images(reader)

        print(f"File        : {path.name}")
        print(f"Size        : {size_kb:.1f} KB ({size_mb:.2f} MB)")
        print(f"Pages       : {len(reader.pages)}")
        print(f"Encrypted   : {'Yes' if reader.is_encrypted else 'No'}")
        print(f"JavaScript  : {'Yes' if javascript else 'No'}")
        print(f"Attachments : {len(attachments)}")
        print(f"Images      : {images}")
        print(f"Fonts       : {len(fonts)}")

        print("\nMetadata:")
        if meta:
            print(f"  Title     : {meta.title or '-'}")
            print(f"  Author    : {meta.author or '-'}")
            print(f"  Subject   : {meta.subject or '-'}")
            print(f"  Creator   : {meta.creator or '-'}")
            print(f"  Producer  : {meta.producer or '-'}")
        else:
            print("  None")

        if fonts:
            print("\nFonts:")

            for font in sorted(fonts):
                print(f"  - {font}")

    except Exception as e:
        print(f"\n[ERROR] Gagal analys file PDF: {e}")